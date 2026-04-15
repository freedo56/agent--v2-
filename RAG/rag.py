import os
import tempfile
import streamlit as st
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_chroma import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import docx
import csv

load_dotenv()

llm = ChatGroq(api_key=os.getenv("GROQ_API_KEY"), model="llama-3.3-70b-versatile")
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")

st.title("📄 RAG — ถามตอบจากเอกสาร")

uploaded_file = st.file_uploader("อัพโหลดไฟล์", type=["pdf", "docx", "txt", "csv"])

def load_file(uploaded_file):
    suffix = "." + uploaded_file.name.split(".")[-1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as f:
        f.write(uploaded_file.read())
        tmp_path = f.name

    if suffix == ".pdf":
        loader = PyPDFLoader(tmp_path)
        return loader.load()
    elif suffix == ".docx":
        doc = docx.Document(tmp_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return [Document(page_content=text)]
    elif suffix == ".txt":
        with open(tmp_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text)]
    elif suffix == ".csv":
        with open(tmp_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            text = "\n".join([", ".join(row) for row in reader])
        return [Document(page_content=text)]

if uploaded_file:
    pages = load_file(uploaded_file)
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = splitter.split_documents(pages)
    vectorstore = Chroma.from_documents(chunks, embeddings)
    st.success(f"โหลดสำเร็จ {len(chunks)} chunks ครับ")

    question = st.text_input("ถามคำถามจากเอกสาร")

    if question:
        docs = vectorstore.similarity_search(question, k=3)
        context = "\n\n".join([d.page_content for d in docs])
        messages = [
            SystemMessage(content=f"ตอบคำถามจากเอกสารนี้เท่านั้น:\n\n{context}"),
            HumanMessage(content=question)
        ]
        response = llm.invoke(messages)
        st.write(response.content)